import os
import sys
import json
import random
import calendar
import datetime
import tempfile
import logging
import threading
from telegram import Update, ReplyKeyboardMarkup, KeyboardButton, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import ApplicationBuilder, CommandHandler, MessageHandler, CallbackQueryHandler, filters, ContextTypes
from telegram.error import TimedOut, NetworkError

_log_dir = os.path.dirname(os.path.abspath(__file__))
_log_file = os.path.join(_log_dir, "bot.log")

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[logging.StreamHandler(), logging.FileHandler(_log_file, encoding="utf-8")],
)
logger = logging.getLogger(__name__)

HAS_DOCX = False
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    logger.warning("python-docx not installed — docx generation disabled")

HAS_NOTION = False
try:
    from notion_client import Client as NotionClient
    NOTION_TOKEN = os.environ.get("NOTION_TOKEN", "")
    NOTION_DB_ID = os.environ.get("NOTION_DB_ID", "")
    if NOTION_TOKEN and NOTION_DB_ID:
        notion_client = NotionClient(auth=NOTION_TOKEN)
        HAS_NOTION = True
        logger.info("Notion integration enabled")
    else:
        logger.warning("NOTION_TOKEN or NOTION_DB_ID not set — Notion disabled")
except ImportError:
    logger.warning("notion-client not installed — Notion disabled")


def save_to_notion(question, answer, user_id, session_type="random"):
    if not HAS_NOTION:
        return
    try:
        notion_client.pages.create(
            parent={"database_id": NOTION_DB_ID},
            properties={
                "Вопрос": {"title": [{"text": {"content": question}}]},
                "Ответ": {"rich_text": [{"text": {"content": answer[:2000]}}]},
                "User ID": {"rich_text": [{"text": {"content": user_id}}]},
                "Дата": {"date": {"start": datetime.datetime.now().strftime("%Y-%m-%d")}},
                "Сессия": {"select": {"name": session_type}},
            },
        )
        logger.info("Saved to Notion: user=%s question=%s", user_id, question[:50])
    except Exception as e:
        logger.error("Failed to save to Notion: %s", e)

_file_lock = threading.Lock()

BOT_TOKEN = os.environ.get("TELEGRAM_BOT_TOKEN", "")
if not BOT_TOKEN:
    logger.error("TELEGRAM_BOT_TOKEN environment variable is not set!")
    sys.exit(1)
DATA_DIR = os.path.dirname(__file__)
QUESTIONS_FILE = os.path.join(DATA_DIR, "questions.json")
ANSWERS_FILE = os.path.join(DATA_DIR, "answers.json")
SELECTED_FILE = os.path.join(DATA_DIR, "selected.json")

DEFAULT_QUESTIONS = [
    "Верю ли я, что моя Высшая Сила может показать мне, как жить и следовать Ее воле?",
    "Вижу ли я в своей сегодняшней жизни старые стереотипы поведения? Если да, то какие именно?",
    "Был ли я обидчивым, эгоистичным, нечестным или трусивым?",
    "Настраивал ли я сам себя на то, чтобы в чем-то разочароваться?",
    "Был ли я добрым и любящим ко всем окружающим?",
    "Беспокоился ли я о вчерашнем или завтрашнем дне?",
    "Позволил ли я себе сегодня впасть в одержимость чем-ли?",
    "Позволил ли я себе быть слишком голодным, злым, одиноким или усталым?",
    "Не отношусь ли я к себе слишком серьезно в какой-то из сфер моей жизни?",
    "Не страдал ли я от физических, психических или духовных проблем?",
    "Не умолчал ли я о том, чем следовало бы поделиться со спонсором?",
    "Испытывал ли я сегодня крайние сильные чувства? Что это были за чувства, и в связи с чем я их испытывал?",
    "Какие сферы моей жизни являются сегодня проблемными?",
    "Какие дефекты проявились сегодня в моей жизни? Как именно?",
    "Был ли сегодня в моей жизни страх?",
    "Что я сегодня сделал такого, чего лучше было бы не делать?",
    "Что я сегодня не сделал из того, что хотел?",
    "Есть ли у меня готовность меняться?",
    "Возникали ли сегодня конфликты в моих отношениях? Какие именно?",
    "Удалось ли мне сохранить верность своим принципам в отношениях с другими людьми?",
    "Не навредил ли я сегодня, прямо или косвенно, себе или кому-либо еще? Как именно?",
    "Не следует ли мне извиниться перед кем-либо или возместить кому-либо ущерб?",
    "В чем я ошибся? Что я сделал бы по-другому, если бы мог? Что мне стоит усвоить на будущее?",
    "Осталась ли я сегодня чистый?",
    "Был ли я сегодня добрым по отношению к самому себе?",
    "Какие чувства я сегодня испытывал? Как они помогли мне поступить в соответствии с духовными принципами?",
    "Как я сегодня служил другим людям?",
    "О каких сегодняшних поступках я вспоминаю с удовольствием?",
    "Что принесло мне сегодня удовлетворение?",
    "Что я сегодня сделал такого, что обязательно нужно повторить?",
    "Был ли я сегодня на собрании? Общался ли с другими выздоравливающими зависимыми?",
    "За что я сегодня благодарен?",
    "Брал ли я сегодня паузу (если да, то сколько раз)?",
    "Получилось (сколько раз) применить практику НВП?",
]

user_states = {}


def _atomic_save(filepath, data):
    with _file_lock:
        dir_name = os.path.dirname(filepath) or "."
        fd, tmp_path = tempfile.mkstemp(dir=dir_name, suffix=".tmp")
        try:
            with os.fdopen(fd, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            os.replace(tmp_path, filepath)
        except Exception:
            try:
                os.remove(tmp_path)
            except OSError:
                pass
            raise


def _safe_load(filepath, default):
    with _file_lock:
        if not os.path.exists(filepath):
            return default() if callable(default) else default
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                return json.load(f)
        except (json.JSONDecodeError, OSError) as e:
            logger.warning("Failed to load %s: %s — using backup", filepath, e)
            backup = filepath + ".bak"
            if os.path.exists(backup):
                try:
                    with open(backup, "r", encoding="utf-8") as f:
                        return json.load(f)
                except Exception:
                    pass
            return default() if callable(default) else default


def _backup_and_save(filepath, data):
    if os.path.exists(filepath):
        try:
            os.replace(filepath, filepath + ".bak")
        except OSError:
            pass
    _atomic_save(filepath, data)


def load_questions():
    return _safe_load(QUESTIONS_FILE, lambda: DEFAULT_QUESTIONS.copy())


def save_questions(questions):
    _backup_and_save(QUESTIONS_FILE, questions)


def load_answers():
    return _safe_load(ANSWERS_FILE, {})


def save_answers(data):
    _backup_and_save(ANSWERS_FILE, data)


def load_selected():
    data = _safe_load(SELECTED_FILE, {})
    result = {}
    for uid, val in data.items():
        if isinstance(val, list):
            result[uid] = {"selected": val, "last_mode": "manual"}
        else:
            result[uid] = val
    return result


def save_selected(data):
    _backup_and_save(SELECTED_FILE, data)


def create_docx(questions, answers, user_id):
    if not HAS_DOCX:
        return None
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    title = doc.add_heading('Инвентаризация', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_str = datetime.datetime.now().strftime("%d.%m.%Y %H:%M")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Дата: {date_str}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    for i, q in enumerate(questions, 1):
        a = answers.get(q, "")
        q_heading = doc.add_heading(f"{i}. {q}", level=2)
        if a:
            p = doc.add_paragraph(a)
            p.paragraph_format.space_after = Pt(12)
        else:
            p = doc.add_paragraph("— не отвечено —")
            p.runs[0].font.color.rgb = RGBColor(150, 150, 150)
            p.paragraph_format.space_after = Pt(12)

    filepath = os.path.join(DATA_DIR, f"inventory_{user_id}.docx")
    doc.save(filepath)
    return filepath


def get_main_keyboard():
    return ReplyKeyboardMarkup(
        [
            [KeyboardButton("🚀 Начать инвентаризацию")],
            [KeyboardButton("🎲 Рандомные вопросы"), KeyboardButton("📋 Выбрать из списка")],
            [KeyboardButton("📖 Ежедневник АН"), KeyboardButton("📊 Создать отчёт")],
        ],
        resize_keyboard=True,
    )


def get_inventory_keyboard():
    return ReplyKeyboardMarkup(
        [[KeyboardButton("⏹ Стоп")]],
        resize_keyboard=True,
    )


MONTHS_RU = [
    "Январь", "Февраль", "Март", "Апрель", "Май", "Июнь",
    "Июль", "Август", "Сентябрь", "Октябрь", "Ноябрь", "Декабрь"
]


def get_calendar_keyboard(year, month, date_from=None, date_to=None):
    buttons = []
    header = "Выбери дату начала:"
    if date_from and not date_to:
        header = f"С: {date_from}\nВыбери дату окончания:"
    elif date_from and date_to:
        header = f"С: {date_from} по: {date_to}"

    buttons.append([
        InlineKeyboardButton("◀️", callback_data=f"cal_prev_{year}_{month}"),
        InlineKeyboardButton(f"{MONTHS_RU[month-1]} {year}", callback_data="cal_ignore"),
        InlineKeyboardButton("▶️", callback_data=f"cal_next_{year}_{month}"),
    ])
    buttons.append([
        InlineKeyboardButton("Пн", callback_data="cal_ignore"),
        InlineKeyboardButton("Вт", callback_data="cal_ignore"),
        InlineKeyboardButton("Ср", callback_data="cal_ignore"),
        InlineKeyboardButton("Чт", callback_data="cal_ignore"),
        InlineKeyboardButton("Пт", callback_data="cal_ignore"),
        InlineKeyboardButton("Сб", callback_data="cal_ignore"),
        InlineKeyboardButton("Вс", callback_data="cal_ignore"),
    ])
    cal = calendar.monthcalendar(year, month)
    for week in cal:
        row = []
        for day in week:
            if day == 0:
                row.append(InlineKeyboardButton(" ", callback_data="cal_ignore"))
            else:
                row.append(InlineKeyboardButton(str(day), callback_data=f"cal_day_{year}_{month}_{day}"))
        buttons.append(row)
    buttons.append([InlineKeyboardButton("⬅️ В меню", callback_data="cal_back")])
    return InlineKeyboardMarkup(buttons), header


def get_selection_keyboard(q_num, total):
    return InlineKeyboardMarkup([
        [
            InlineKeyboardButton("✅ Выбрать", callback_data=f"sel_{q_num}"),
            InlineKeyboardButton("⏭ Пропустить", callback_data=f"skip_{q_num}"),
        ],
        [
            InlineKeyboardButton("⬅️ В меню", callback_data="cancel_select"),
        ],
    ])


def get_summary_keyboard(selected_count):
    return InlineKeyboardMarkup([
        [InlineKeyboardButton(f"🚀 Начать инвентаризацию ({selected_count})", callback_data="start_from_select")],
        [InlineKeyboardButton("➕ Добавить свой вопрос", callback_data="add_own_question")],
        [InlineKeyboardButton("🔄 Выбрать заново", callback_data="restart_select")],
        [InlineKeyboardButton("⬅️ В меню", callback_data="cancel_select")],
    ])


async def _safe_reply(message, text, **kwargs):
    for attempt in range(3):
        try:
            return await message.reply_text(text, **kwargs)
        except TimedOut:
            logger.warning("TimedOut on attempt %d, retrying...", attempt + 1)
            if attempt == 2:
                raise
        except NetworkError as e:
            logger.warning("NetworkError on attempt %d: %s", attempt + 1, e)
            if attempt == 2:
                raise


async def _safe_edit(query, text, **kwargs):
    try:
        return await query.edit_message_text(text, **kwargs)
    except Exception as e:
        logger.warning("edit_message_text failed: %s", e)
        return None


async def error_handler(update, context):
    logger.error("Exception while handling an update: %s", context.error, exc_info=context.error)
    if isinstance(context.error, TimedOut):
        logger.warning("TimedOut error — retrying next update")
        return
    if update and hasattr(update, "callback_query") and update.callback_query:
        try:
            await update.callback_query.message.reply_text(
                "Произошла ошибка. Попробуйте ещё раз.",
                reply_markup=get_main_keyboard(),
            )
        except Exception:
            pass
    elif update and hasattr(update, "message") and update.message:
        try:
            await update.message.reply_text(
                "Произошла ошибка. Попробуйте ещё раз.",
                reply_markup=get_main_keyboard(),
            )
        except Exception:
            pass


async def cmd_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = str(update.effective_user.id)
    user_states[user_id] = {}
    await update.message.reply_text(
        "Привет! Я бот для инвентаризации.\nВыбери действие:",
        reply_markup=get_main_keyboard(),
    )


async def cmd_stop(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = str(update.effective_user.id)
    user_states[user_id] = {}
    await update.message.reply_text(
        "Инвентаризация остановлена.",
        reply_markup=get_main_keyboard(),
    )


async def handle_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    user_id = str(query.from_user.id)
    data = query.data
    questions = load_questions()
    all_sel = load_selected()

    if user_id not in user_states:
        saved = all_sel.get(user_id, {})
        user_states[user_id] = {"selected": saved.get("selected", []), "last_mode": saved.get("last_mode", "random")}

    state = user_states[user_id]

    try:
        if data.startswith("sel_") or data.startswith("skip_"):
            idx = int(data.split("_")[1])
            is_selected = data.startswith("sel_")

            if is_selected:
                if idx not in state["selected"]:
                    state["selected"].append(idx)

            all_sel[user_id] = {"selected": state["selected"], "last_mode": state.get("last_mode", "manual")}
            save_selected(all_sel)

            next_idx = idx + 1
            if next_idx >= len(questions):
                selected = state["selected"]
                if selected:
                    text = f"Выбрано вопросов: {len(selected)}\n\n"
                    for i, qi in enumerate(selected, 1):
                        text += f"{i}. {questions[qi]}\n"
                    if len(text) > 4000:
                        text = text[:4000] + "..."
                    await _safe_edit(query, text, reply_markup=get_summary_keyboard(len(selected)))
                else:
                    await _safe_edit(query, "Ни один вопрос не выбран.", reply_markup=get_summary_keyboard(0))
            else:
                await _safe_edit(
                    query,
                    f"({next_idx + 1}/{len(questions)})\n\n{questions[next_idx]}",
                    reply_markup=get_selection_keyboard(next_idx, len(questions)),
                )
            return

        elif data == "start_from_select":
            selected = state.get("selected", [])
            if not selected:
                await _safe_edit(query, "Ни один вопрос не выбран!", reply_markup=get_summary_keyboard(0))
                return
            state["last_mode"] = "manual"
            all_sel[user_id] = {"selected": selected, "last_mode": "manual"}
            save_selected(all_sel)
            user_states[user_id] = {
                "collecting": True,
                "step": 0,
                "selected_list": selected,
                "last_mode": "manual",
            }
            await _safe_edit(query, f"Инвентаризация ({len(selected)} из списка)!\n\nВопрос 1/{len(selected)}:\n\n{questions[selected[0]]}")
            await _safe_reply(query.message, "Отвечай текстом:", reply_markup=get_inventory_keyboard())
            return

        elif data == "cancel_select":
            user_states[user_id] = {}
            await _safe_edit(query, "Выбери действие:")
            await _safe_reply(query.message, "Выбери действие:", reply_markup=get_main_keyboard())
            return

        elif data.startswith("cal_prev_") or data.startswith("cal_next_"):
            parts = data.split("_")
            year = int(parts[2])
            month = int(parts[3])
            if data.startswith("cal_prev_"):
                month -= 1
                if month < 1:
                    month = 12
                    year -= 1
            else:
                month += 1
                if month > 12:
                    month = 1
                    year += 1
            date_from = state.get("cal_from")
            date_to = state.get("cal_to")
            keyboard, header = get_calendar_keyboard(year, month, date_from, date_to)
            await _safe_edit(query, header, reply_markup=keyboard)

        elif data.startswith("cal_day_"):
            parts = data.split("_")
            year = int(parts[2])
            month = int(parts[3])
            day = int(parts[4])
            date_str = f"{day:02d}.{month:02d}.{year}"

            if "cal_from" not in state:
                state["cal_from"] = date_str
                state["cal_year"] = year
                state["cal_month"] = month
                user_states[user_id] = state
                keyboard, header = get_calendar_keyboard(year, month, date_str)
                await _safe_edit(query, header, reply_markup=keyboard)
            else:
                date_from = state["cal_from"]
                d_from = datetime.datetime.strptime(date_from, "%d.%m.%Y")
                d_to = datetime.datetime.strptime(date_str, "%d.%m.%Y")
                if d_to < d_from:
                    date_from, date_str = date_str, date_from
                    d_from, d_to = d_to, d_from
                state["cal_to"] = date_str
                user_states[user_id] = state

                all_ans = load_answers()
                user_answers = all_ans.get(user_id, {})
                questions_list = load_questions()
                answered = []
                for i, q in enumerate(questions_list, 1):
                    a = user_answers.get(q, "")
                    if a:
                        answered.append((i, q, a))

                if HAS_DOCX:
                    doc = Document()
                    style = doc.styles['Normal']
                    font = style.font
                    font.name = 'Arial'
                    font.size = Pt(11)

                    title = doc.add_heading('Инвентаризация', level=0)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(f"Период: с {date_from} по {date_str}")
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(100, 100, 100)

                    doc.add_paragraph()

                    if answered:
                        for i, q, a in answered:
                            doc.add_heading(f"{i}. {q}", level=2)
                            p = doc.add_paragraph(a)
                            p.paragraph_format.space_after = Pt(12)
                    else:
                        doc.add_paragraph("Нет ответов за выбранный период.")

                    filepath = os.path.join(DATA_DIR, f"report_{user_id}.docx")
                    doc.save(filepath)

                    await _safe_edit(query, f"Отчёт за период с {date_from} по {date_str}\n({len(answered)} вопросов)")
                    await _safe_reply(query.message, "Выбери действие:", reply_markup=get_main_keyboard())
                    with open(filepath, "rb") as f:
                        await query.message.reply_document(
                            document=f,
                            filename="Отчёт.docx",
                            caption="Отчёт готов",
                        )
                    os.remove(filepath)
                else:
                    lines = [f"Отчёт за период с {date_from} по {date_str}", ""]
                    for i, q, a in answered:
                        lines.append(f"{i}. {q}\n{a}\n")
                    if not answered:
                        lines.append("Нет ответов за выбранный период.")
                    await _safe_edit(query, "\n".join(lines)[:4000])
                state.pop("cal_from", None)
                state.pop("cal_to", None)
                user_states[user_id] = state
                await _safe_reply(query.message, "Выбери действие:", reply_markup=get_main_keyboard())

        elif data == "cal_back":
            user_states[user_id] = {}
            await _safe_edit(query, "Выбери действие:")
            await _safe_reply(query.message, "Выбери действие:", reply_markup=get_main_keyboard())

        elif data == "cal_ignore":
            await query.answer()

        elif data == "add_own_question":
            user_states[user_id] = {"adding_own": True, "selected": state.get("selected", [])}
            await _safe_edit(query, "Напиши свой вопрос:")
            return

        elif data == "restart_select":
            user_states[user_id] = {"selected": [], "last_mode": "manual"}
            await _safe_edit(
                query,
                f"Вопрос 1/{len(questions)}:\n\n{questions[0]}",
                reply_markup=get_selection_keyboard(0, len(questions)),
            )
            return

    except Exception as e:
        logger.error("Error in handle_callback: %s", e, exc_info=True)
        try:
            await query.message.reply_text("Произошла ошибка. Выбери действие:", reply_markup=get_main_keyboard())
        except Exception:
            pass


MAIN_BUTTONS = {"🚀 Начать инвентаризацию", "🎲 Рандомные вопросы", "📋 Выбрать из списка", "📖 Ежедневник АН", "📊 Создать отчёт", "⏹ Стоп"}


async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = str(update.effective_user.id)
    text = update.message.text.strip()

    if user_id not in user_states:
        user_states[user_id] = {}

    if text in MAIN_BUTTONS:
        if user_states[user_id].get("collecting"):
            user_states[user_id] = {}
            await _safe_reply(update.message, "Инвентаризация остановлена.", reply_markup=get_main_keyboard())
            return

    all_answers = load_answers()
    questions = load_questions()
    all_sel = load_selected()

    if user_id not in all_answers:
        all_answers[user_id] = {}

    state = user_states[user_id]
    if "selected" not in state:
        saved = all_sel.get(user_id, {})
        state["selected"] = saved.get("selected", [])
        state["last_mode"] = saved.get("last_mode", "random")

    if text == "⏹ Стоп":
        user_states[user_id] = {}
        await _safe_reply(update.message, "Инвентаризация остановлена.", reply_markup=get_main_keyboard())
        return

    if state.get("adding_own"):
        questions.append(text)
        save_questions(questions)
        new_idx = len(questions) - 1
        state["adding_own"] = False
        if new_idx not in state.get("selected", []):
            state.setdefault("selected", []).append(new_idx)
        user_states[user_id] = state
        selected = state["selected"]
        if selected:
            summary = f"Добавлен!\n\nВыбрано вопросов: {len(selected)}\n\n"
            for i, qi in enumerate(selected, 1):
                summary += f"{i}. {questions[qi]}\n"
            if len(summary) > 4000:
                summary = summary[:4000] + "..."
            await _safe_reply(update.message, summary, reply_markup=get_summary_keyboard(len(selected)))
        else:
            await _safe_reply(update.message, f"Добавлен: {text}", reply_markup=get_main_keyboard())
        return

    if state.get("collecting"):
        selected_list = state.get("selected_list", [])
        step = state.get("step", 0)

        if not selected_list:
            user_states[user_id] = {}
            await _safe_reply(update.message, "Нет выбранных вопросов.", reply_markup=get_main_keyboard())
            return

        valid_list = [i for i in selected_list if i < len(questions)]
        if not valid_list:
            user_states[user_id] = {}
            await _safe_reply(update.message, "Нет доступных вопросов.", reply_markup=get_main_keyboard())
            return
        state["selected_list"] = valid_list
        user_states[user_id] = state
        selected_list = valid_list

        if step >= len(selected_list):
            user_states[user_id] = {}
            selected_questions = [questions[i] for i in selected_list]
            filepath = create_docx(selected_questions, all_answers.get(user_id, {}), user_id)
            if filepath:
                with open(filepath, "rb") as f:
                    await update.message.reply_document(
                        document=f,
                        filename="Инвентаризация.docx",
                        caption="Инвентаризация завершена",
                    )
                os.remove(filepath)
            else:
                await _safe_reply(update.message, "Инвентаризация завершена (python-docx не установлен).")
            await _safe_reply(update.message, "Выбери действие:", reply_markup=get_main_keyboard())
            return

        q_idx = selected_list[step]
        if q_idx >= len(questions):
            step += 1
            state["step"] = step
            user_states[user_id] = state
            await _safe_reply(
                update.message,
                f"Вопрос {step}/{len(selected_list)} не найден в списке, пропускаю...",
                reply_markup=get_inventory_keyboard(),
            )
            return
        current_q = questions[q_idx]
        all_answers[user_id][current_q] = text
        save_answers(all_answers)
        session_type = state.get("last_mode", "random")
        save_to_notion(current_q, text, user_id, session_type)
        step += 1
        state["step"] = step
        user_states[user_id] = state

        if step >= len(selected_list):
            user_states[user_id] = {}
            selected_questions = [questions[i] for i in selected_list]
            filepath = create_docx(selected_questions, all_answers.get(user_id, {}), user_id)
            if filepath:
                with open(filepath, "rb") as f:
                    await update.message.reply_document(
                        document=f,
                        filename="Инвентаризация.docx",
                        caption="Инвентаризация завершена",
                    )
                os.remove(filepath)
            else:
                await _safe_reply(update.message, "Инвентаризация завершена (python-docx не установлен).")
            await _safe_reply(update.message, "Выбери действие:", reply_markup=get_main_keyboard())
        else:
            next_q_idx = selected_list[step]
            if next_q_idx >= len(questions):
                step += 1
                state["step"] = step
                user_states[user_id] = state
                return
            await _safe_reply(
                update.message,
                f"[{step}/{len(selected_list)}]\n\n{questions[next_q_idx]}",
                reply_markup=get_inventory_keyboard(),
            )
        return

    if text == "🚀 Начать инвентаризацию":
        saved = all_sel.get(user_id, {})
        saved_list = saved.get("selected", [])
        last_mode = state.get("last_mode", saved.get("last_mode", "random"))
        if last_mode == "manual" and saved_list:
            selected = saved_list
            mode_text = f"выбрано из списка: {len(selected)}"
        else:
            fixed = [len(questions) - 3, len(questions) - 2, len(questions) - 1]
            pool = [i for i in range(len(questions)) if i not in fixed]
            random_count = min(5, len(pool))
            random_part = random.sample(pool, random_count)
            selected = random_part + fixed
            mode_text = f"{random_count} случайных + 3 обязательных"
        user_states[user_id] = {
            "collecting": True,
            "step": 0,
            "selected_list": selected,
            "last_mode": last_mode,
        }
        await _safe_reply(
            update.message,
            f"Инвентаризация ({mode_text})!\n\nВопрос 1/{len(selected)}:\n\n{questions[selected[0]]}",
            reply_markup=get_inventory_keyboard(),
        )
        return

    if text == "🎲 Рандомные вопросы":
        fixed = [len(questions) - 3, len(questions) - 2, len(questions) - 1]
        pool = [i for i in range(len(questions)) if i not in fixed]
        random_count = min(5, len(pool))
        random_part = random.sample(pool, random_count)
        selected = random_part + fixed
        user_states[user_id] = {
            "collecting": True,
            "step": 0,
            "selected_list": selected,
            "last_mode": "random",
        }
        await _safe_reply(
            update.message,
            f"Случайные вопросы ({random_count} + 3 обязательных)!\n\nВопрос 1/{len(selected)}:\n\n{questions[selected[0]]}",
            reply_markup=get_inventory_keyboard(),
        )
        return

    if text == "📋 Выбрать из списка":
        state["selected"] = []
        state["last_mode"] = "manual"
        user_states[user_id] = state
        await _safe_reply(
            update.message,
            f"Вопрос 1/{len(questions)}:\n\n{questions[0]}",
            reply_markup=get_selection_keyboard(0, len(questions)),
        )
        return

    if text == "📖 Ежедневник АН":
        keyboard = InlineKeyboardMarkup([
            [InlineKeyboardButton("Открыть Ежедневник АН", url="https://na-russia.org/meditation-today")],
        ])
        await _safe_reply(update.message, "Нажми кнопку чтобы открыть:", reply_markup=keyboard)
        return

    if text == "📊 Создать отчёт":
        now = datetime.datetime.now()
        state["cal_year"] = now.year
        state["cal_month"] = now.month
        user_states[user_id] = state
        keyboard, header = get_calendar_keyboard(now.year, now.month)
        await _safe_reply(update.message, header, reply_markup=keyboard)
        return

    await _safe_reply(update.message, "Выбери действие:", reply_markup=get_main_keyboard())


def main():
    pid_file = os.path.join(DATA_DIR, "bot.pid")

    if os.path.exists(pid_file):
        try:
            with open(pid_file, "r") as f:
                old_pid = int(f.read().strip())
            if sys.platform == "win32":
                import ctypes
                kernel32 = ctypes.windll.kernel32
                handle = kernel32.OpenProcess(0x1000, False, old_pid)
                if handle:
                    kernel32.CloseHandle(handle)
                    logger.warning("Previous bot process (PID %d) still alive. Killing it.", old_pid)
                    os.system(f"taskkill /PID {old_pid} /F >nul 2>&1")
            else:
                os.kill(old_pid, 0)
                logger.warning("Previous bot process (PID %d) still alive. Killing it.", old_pid)
                os.kill(old_pid, 9)
        except (OSError, ProcessLookupError):
            pass

    with open(pid_file, "w") as f:
        f.write(str(os.getpid()))

    logger.info("Bot starting (PID %d)...", os.getpid())
    try:
        app = (
            ApplicationBuilder()
            .token(BOT_TOKEN)
            .connect_timeout(30)
            .read_timeout(30)
            .write_timeout(30)
            .build()
        )
        app.add_handler(CommandHandler("start", cmd_start))
        app.add_handler(CommandHandler("stop", cmd_stop))
        app.add_handler(CallbackQueryHandler(handle_callback))
        app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))
        app.add_error_handler(error_handler)
        app.run_polling(drop_pending_updates=True)
    except Exception as e:
        logger.exception("Bot crashed: %s", e)
        raise
    finally:
        if os.path.exists(pid_file):
            try:
                os.remove(pid_file)
            except OSError:
                pass
        logger.info("Bot stopped.")


if __name__ == "__main__":
    main()
